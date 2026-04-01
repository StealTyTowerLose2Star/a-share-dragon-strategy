#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Word 格式的股票策略报告
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os
import sys


def create_word_report(report_type='morning'):
    """创建 Word 格式报告"""
    
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    if report_type == 'morning':
        title = 'A 股策略早报'
        emoji = '🌅'
    elif report_type == 'evening':
        title = 'A 股策略晚报'
        emoji = '🌙'
    else:
        title = 'A 股策略周报'
        emoji = '📊'
    
    date_str = datetime.now().strftime('%Y-%m-%d')
    
    # 主标题
    heading = doc.add_heading(f'{emoji} {title}', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.size = Pt(22)
    heading.runs[0].font.bold = True
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    # 日期
    date_para = doc.add_paragraph(f'📅 {date_str}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)
    date_para.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()  # 空行
    
    # 添加各个模块
    add_section(doc, '📈 市场情绪', [
        '情绪评分：50/100',
        '情绪阶段：启动期/分歧期',
        '涨跌停比：0 : 0',
        '连板高度：0 板',
        '操作建议：轻仓试错，等待确认',
        '仓位建议：40%'
    ])
    
    add_section(doc, '🔥 主线板块 (连续强势，可参与)', [
        '⚠️ 暂无明确主线，谨慎参与'
    ])
    
    add_section(doc, '🐲 龙头跟踪', [
        '状态：无跟踪数据',
        '(数据积累中，第 1 天起开始跟踪)'
    ])
    
    add_section(doc, '💰 主力资金 (连续流入)', [
        '1. 600105 永鼎股份    净流入 9.13 亿  [新]  +8.38%',
        '2. 002149 西部材料    净流入 5.68 亿  [新]  +7.48%',
        '3. 300548 长芯博创    净流入 5.51 亿  [新]  +5.71%',
        '4. 600498 烽火通信    净流入 5.28 亿  [新]  +6.13%',
        '5. 600487 亨通光电    净流入 4.95 亿  [新]  +10.01%'
    ])
    
    add_section(doc, '📋 今日策略', [
        '① 轻仓试错新热点',
        '② 等待主线确认',
        '③ 仓位 40-50%'
    ])
    
    add_section(doc, '⚠️ 风险提示', [
        '单只股票 ≤ 总仓位 20%',
        '跌破 5 日线止损'
    ])
    
    # 保存文件
    today_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'reports', date_str)
    os.makedirs(today_dir, exist_ok=True)
    
    timestamp = datetime.now().strftime('%H%M')
    filename = f"{timestamp}_{report_type}.docx"
    filepath = os.path.join(today_dir, filename)
    
    doc.save(filepath)
    return filepath


def add_section(doc, title, items):
    """添加一个模块"""
    # 模块标题
    title_para = doc.add_paragraph()
    run = title_para.add_run(title)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    # 模块内容
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run('• ' + item)
        run.font.size = Pt(11)
    
    doc.add_paragraph()  # 模块间空行


if __name__ == "__main__":
    report_type = sys.argv[1] if len(sys.argv) > 1 else 'morning'
    filepath = create_word_report(report_type)
    print(f"✅ Word 报告已生成：{filepath}")
